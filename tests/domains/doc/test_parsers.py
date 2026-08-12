"""M3-01 golden 测试：txt/md/pdf/docx 解析与损坏文件友好报错。"""

from __future__ import annotations

from pathlib import Path

import pytest

from littledotmcp.rag.parsers import ParseError, parse_document


def _write(path: Path, data: bytes | str) -> Path:
    path.write_bytes(data if isinstance(data, bytes) else data.encode("utf-8"))
    return path


def test_parse_txt(tmp_path: Path) -> None:
    p = _write(tmp_path / "tmp.txt", "第一行\n第二行")
    parsed = parse_document(p)
    assert parsed.mime == "text/plain"
    assert parsed.text == "第一行\n第二行"
    assert parsed.char_count == len("第一行\n第二行")
    assert parsed.page_count == 0


def test_parse_md(tmp_path: Path) -> None:
    p = _write(tmp_path / "tmp.md", "# 标题\n\n- 列表项")
    parsed = parse_document(p)
    assert parsed.mime == "text/markdown"
    assert "# 标题" in parsed.text and "列表项" in parsed.text


def test_parse_pdf(tmp_path: Path) -> None:
    from pypdf import PdfWriter

    target = tmp_path / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    with target.open("wb") as f:
        writer.write(f)

    parsed = parse_document(target)
    assert parsed.mime == "application/pdf"
    assert parsed.page_count == 2
    assert isinstance(parsed.text, str)


def test_parse_docx(tmp_path: Path) -> None:
    import docx

    target = tmp_path / "sample.docx"
    document = docx.Document()
    document.add_paragraph("标题段落")
    document.add_paragraph("内容段落二")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"
    document.save(str(target))

    parsed = parse_document(target)
    assert parsed.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "标题段落" in parsed.text and "内容段落二" in parsed.text
    assert "a | b" in parsed.text
    assert parsed.page_count == 0


def test_unsupported_extension(tmp_path: Path) -> None:
    p = _write(tmp_path / "x.xyz", "hello")
    with pytest.raises(ParseError) as exc:
        parse_document(p)
    assert "不支持" in str(exc.value)


def test_missing_file(tmp_path: Path) -> None:
    with pytest.raises(ParseError) as exc:
        parse_document(tmp_path / "nope.txt")
    assert "不存在" in str(exc.value)


def test_corrupted_pdf(tmp_path: Path) -> None:
    p = _write(tmp_path / "bad.pdf", b"%PDF-1.4 not really a pdf ........")
    with pytest.raises(ParseError):
        parse_document(p)


def test_corrupted_docx(tmp_path: Path) -> None:
    p = _write(tmp_path / "bad.docx", b"PK\x03\x04 broken docx payload")
    with pytest.raises(ParseError):
        parse_document(p)


def test_oversized_text_truncated(tmp_path: Path) -> None:
    p = _write(tmp_path / "big.txt", "x" * 3_000_000)
    parsed = parse_document(p)
    assert parsed.char_count == 2_000_000
    assert len(parsed.text) == 2_000_000
