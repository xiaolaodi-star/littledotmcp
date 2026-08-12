"""M3-01 文档解析器集合：txt/md/pdf/docx → 纯文本 + 元信息。

- txt/md：原生按 UTF-8 读取（容错 replace，坏字节不崩溃）
- pdf：优先 pypdf，失败回退 pdfplumber；返回页数
- docx：python-docx 提取段落与表格文本；页数不可得置 0
- 解析失败统一抛 ParseError（可读信息，不泄露堆栈）
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from ..common.errors import ValidationError
from ..common.logging import get_logger

logger = get_logger(__name__)

# 单文件解析文本上限（防超大文件打爆内存；超限截断并告警）
_MAX_CHARS = 2_000_000

# 扩展名 → MIME
_EXT_MIME = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class ParseError(ValidationError):
    """文档解析失败（含不支持的格式/损坏文件）。"""


@dataclass
class ParsedDoc:
    """解析结果：纯文本 + 元信息。"""

    text: str
    mime: str
    page_count: int
    char_count: int


def parse_document(path: Path) -> ParsedDoc:
    """按扩展名路由解析文档，返回纯文本与元信息。

    Args:
        path: 文档路径（扩展名决定解析器）。

    Returns:
        ParsedDoc 含 text/mime/page_count/char_count。

    Raises:
        ParseError: 文件不存在、格式不支持或内容损坏。
    """
    suffix = path.suffix.lower()
    if suffix not in _EXT_MIME:
        raise ParseError(
            f"不支持的文档类型 {suffix or '(无扩展名)'!r}，仅支持：{', '.join(sorted(_EXT_MIME))}"
        )
    if not path.is_file():
        raise ParseError(f"文件不存在：{path}")

    try:
        if suffix in (".txt", ".md"):
            parsed = _parse_text(path, _EXT_MIME[suffix])
        elif suffix == ".pdf":
            parsed = _parse_pdf(path)
        else:
            parsed = _parse_docx(path)
    except ParseError:
        raise
    except Exception as exc:  # 解析库内部异常统一转可读错误
        raise ParseError(f"解析失败：{path.name}：{exc}") from exc

    if parsed.char_count > _MAX_CHARS:
        logger.warning("文档 %s 超过 %d 字符，已截断", path.name, _MAX_CHARS)
        return ParsedDoc(
            text=parsed.text[:_MAX_CHARS],
            mime=parsed.mime,
            page_count=parsed.page_count,
            char_count=_MAX_CHARS,
        )
    return parsed


def _parse_text(path: Path, mime: str) -> ParsedDoc:
    text = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDoc(text=text, mime=mime, page_count=0, char_count=len(text))


def _parse_pdf(path: Path) -> ParsedDoc:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ParseError("缺少 pypdf 依赖，请执行 uv add pypdf") from exc

    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages)
        return ParsedDoc(
            text=text, mime=_EXT_MIME[".pdf"], page_count=len(reader.pages), char_count=len(text)
        )
    except ParseError:
        raise
    except Exception as exc:  # pypdf 失败回退 pdfplumber
        logger.warning("pypdf 解析失败（%s），回退 pdfplumber", exc)
        return _parse_pdf_with_pdfplumber(path)


def _parse_pdf_with_pdfplumber(path: Path) -> ParsedDoc:
    try:
        import pdfplumber
    except ImportError as exc:
        raise ParseError("缺少 pdfplumber 依赖，请执行 uv add pdfplumber") from exc

    try:
        with pdfplumber.open(str(path)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages)
            return ParsedDoc(
                text=text, mime=_EXT_MIME[".pdf"], page_count=len(pdf.pages), char_count=len(text)
            )
    except ParseError:
        raise
    except Exception as exc:
        raise ParseError(f"PDF 解析失败（pypdf/pdfplumber 均不可用）：{exc}") from exc


def _parse_docx(path: Path) -> ParsedDoc:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise ParseError("缺少 python-docx 依赖，请执行 uv add python-docx") from exc

    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    # 表格内容逐行拼接（单元格以 | 分隔）
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return ParsedDoc(text=text, mime=_EXT_MIME[".docx"], page_count=0, char_count=len(text))
